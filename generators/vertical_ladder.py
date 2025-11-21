"""
Модуль генерации Word-документов
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path
from logger import app_logger
import config
from generators.base_generator import BaseProtocolGenerator
from validator import DataValidator


class VerticalLadderGenerator(BaseProtocolGenerator):
    """Генератор протоколов для вертикальных лестниц"""
    
    def __init__(self, data):
        super().__init__(data)
        self.document = None
    
    def validate(self):
        is_valid, errors = DataValidator.validate_all_data(self.data)
        if not is_valid:
            raise ValueError("\n".join(errors))
    
    def generate_doc(self, output_path=None):
        """
        Создает Word-документ на основе переданных данных
        
        Args:
            data (dict): Словарь с данными для документа
                - date: дата
                - customer: заказчик
                - object_full_address: адрес/наименование испытываемого объекта (из п.1.2 договора)
                - ladders: список лестниц, каждая со своими характеристиками:
                    [
                        {
                            'number': int,
                            'name': str,
                            'height': str,
                            'width': str,
                            'steps_count': str,
                            'mount_points': str,
                            'platform_length': str,
                            'platform_width': str,
                            'fence_height': str,
                            'wall_distance': str,
                            'ground_distance': str,
                            'step_distance': str
                        },
                        ...
                    ]
                - test_time: время проведения испытаний
                - temperature: температура воздуха
                - wind_speed: скорость ветра
                - damage_found: внешние повреждения обнаружены (bool)
                - mount_violation_found: нарушение крепления обнаружено (bool)
                - weld_violation_found: нарушение сварных швов обнаружено (bool)
                - paint_compliant: соответствие окраски ГОСТ 9.302 (bool)
        
        Таблица результатов испытаний формируется автоматически на основе типа лестницы:
        - П1-1 (≤6м): ограждения площадки - 2 точки
        - П1-2 (>6м): ограждения лестницы и площадки - 2 + (высота / 1.2) точек
        
        Returns:
            str: Путь к созданному файлу
        """
        try:
            data = self.data
            app_logger.info(f"Начало генерации документа для заказчика: {data.get('customer')}")
            
            self._set_document()
            
            # Шапка с логотипом и реквизитами
            self._add_company_header()
            
            # Заголовок документа
            self._add_header(data)
            
            # Основная информация
            self._add_main_info(data)
            
            # Характеристики объекта
            self._add_object_characteristics(data)
            
            # Условия проведения испытаний
            self._add_test_conditions(data)
            
            # Средства проведения испытаний
            self._add_test_equipment(data)
            
            # Визуальный осмотр
            self._add_visual_inspection(data)
            
            # Расчет величины нагрузки
            self._add_load_calculation(data)
            
            # Испытаниям подлежат
            # self._add_test_subjects(data)  # Убрано из приложения
            
            # Автоматическая таблица результатов испытаний
            self._add_dynamic_table(data)
            
            # Выводы по результатам испытаний
            self._add_conclusions(data)
            
            # Подписи и дата
            self._add_signatures(data)
            
            # Сохранение файла
            filename = self._generate_filename(data)
            filepath = self._resolve_output_path(output_path, filename)
            self.document.save(str(filepath))
            
            app_logger.info(f"Документ успешно создан: {filepath}")
            return str(filepath)
            
        except Exception as e:
            app_logger.error(f"Ошибка при создании документа: {e}")
            raise
    
    def _setup_styles(self):
        """Настройка стилей документа"""
        styles = self.document.styles
        
        # Стиль Normal (базовый для всего документа)
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(10)
        normal_style.font.color.rgb = RGBColor(0, 0, 0)
        
        # Настройка форматирования абзаца для стиля Normal
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_style.paragraph_format.left_indent = Pt(0)
        normal_style.paragraph_format.right_indent = Pt(0)
        normal_style.paragraph_format.first_line_indent = None
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.space_after = Pt(0)
        normal_style.paragraph_format.line_spacing = 1.0
        
        # Стиль для заголовков
        if 'CustomHeading' not in styles:
            style = styles.add_style('CustomHeading', 1)
            style.font.name = 'Times New Roman'
            style.font.size = Pt(10)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
    
    def _format_paragraph(self, paragraph):
        """Применяет стандартное форматирование к абзацу согласно настройкам"""
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
    
    def _add_company_header(self):
        """Добавляет шапку с логотипом и реквизитами компании в виде таблицы"""
        # Таблица для размещения логотипа и реквизитов (без границ)
        header_table = self.document.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.allow_autofit = False
        
        # Убираем границы таблицы
        for row in header_table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(
                    OxmlElement('w:tcBorders')
                )
        
        # Левая ячейка - логотип
        logo_cell = header_table.rows[0].cells[0]
        logo_cell.width = Inches(1.5)
        logo_cell.vertical_alignment = 1  # Выравнивание по центру вертикально
        
        # Попытка добавить логотип - заново проверяем путь
        logo_file = config.get_logo_file()
        app_logger.info(f"Проверка логотипа: {logo_file}, существует: {logo_file.exists() if logo_file else False}")
        
        if logo_file and logo_file.exists():
            try:
                logo_paragraph = logo_cell.paragraphs[0]
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(str(logo_file.absolute()), height=Inches(0.8))
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                logo_paragraph.paragraph_format.space_before = Pt(0)
                logo_paragraph.paragraph_format.space_after = Pt(0)
                app_logger.info(f"Логотип успешно добавлен: {logo_file}")
            except Exception as e:
                app_logger.error(f"Не удалось добавить логотип: {e}")
                logo_cell.text = config.COMPANY_NAME
                para = logo_cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
        else:
            app_logger.warning(f"Логотип не найден: {logo_file}")
            logo_cell.text = config.COMPANY_NAME
            para = logo_cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
        
        # Правая ячейка - реквизиты простым текстом
        details_cell = header_table.rows[0].cells[1]
        details_cell.width = Inches(4.0)
        details_cell.vertical_alignment = 1  # Выравнивание по центру вертикально
        
        # Удаляем стандартный параграф
        details_cell.text = ''
        
        # Название компании (жирным)
        p = details_cell.paragraphs[0]
        run = p.add_run(config.COMPANY_NAME)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        # Адрес строка 1
        p = details_cell.add_paragraph()
        run = p.add_run(config.COMPANY_ADDRESS_LINE1)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        # Адрес строка 2
        p = details_cell.add_paragraph()
        run = p.add_run(config.COMPANY_ADDRESS_LINE2)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        # Телефон
        p = details_cell.add_paragraph()
        run = p.add_run(config.COMPANY_PHONE)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        # Email
        p = details_cell.add_paragraph()
        run = p.add_run(config.COMPANY_EMAIL)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        # Сайт
        p = details_cell.add_paragraph()
        run = p.add_run(config.COMPANY_WEBSITE)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        # Отступ после шапки
        p = self.document.add_paragraph()
        self._format_paragraph(p)
    
    def _add_header(self, data):
        """Добавляет заголовок документа"""
        title = 'Протокол испытания вертикальных пожарных лестниц'
        
        heading = self.document.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Подзаголовок с датой
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.paragraph_format.space_before = Pt(0)
        date_para.paragraph_format.space_after = Pt(0)
        date_para.paragraph_format.line_spacing = 1.0
        date_run = date_para.add_run(f"от {data.get('date', '')}")
        date_run.font.size = Pt(10)
        date_run.font.bold = True
        date_run.font.name = 'Times New Roman'
        date_run.font.color.rgb = RGBColor(0, 0, 0)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)
    
    def _add_main_info(self, data):
        """Добавляет основную информацию"""
        # Заказчик
        p = self.document.add_paragraph()
        run = p.add_run('Заказчик: ')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p.add_run(data.get('customer', ''))
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        # Адрес/наименование испытываемого объекта (объединённое поле)
        p = self.document.add_paragraph()
        run = p.add_run('Адрес/наименование испытываемого объекта: ')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p.add_run(data.get('object_full_address', ''))
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)
    
    def _add_object_characteristics(self, data):
        """Добавляет характеристику испытываемых конструкций"""
        heading = self.document.add_heading('Характеристика испытываемых конструкций', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Получаем список лестниц
        ladders = data.get('ladders', [])
        
        # Если нет лестниц - создаём одну по умолчанию из старых данных
        if not ladders:
            ladders = [{
                'number': 1,
                'name': '',
                'height': data.get('ladder_height', ''),
                'width': data.get('ladder_width', ''),
                'steps_count': data.get('steps_count', ''),
            }]
        
        # Выводим информацию по каждой лестнице
        for ladder in ladders:
            ladder_num = ladder.get('number', 1)
            ladder_name = ladder.get('name', '')
            ladder_internal_type = ladder.get('ladder_type', 'vertical')
            if ladder_internal_type not in ('', 'vertical'):
                app_logger.warning(f"Лестница №{ladder_num}: получен устаревший тип '{ladder_internal_type}', используется вертикальная схема описания")
            
            # Определяем тип по высоте
            try:
                ladder_height_value = float(str(ladder.get('height', '0')).replace(',', '.'))
                ladder_type = 'П1-1' if ladder_height_value <= 6 else 'П1-2'
            except (ValueError, TypeError):
                ladder_type = ''
            
            # Данные лестницы - через запятую в строку (только заполненные поля)
            characteristics_parts = []
            
            # Тип лестницы всегда выводим
            if ladder_type:
                characteristics_parts.append(f"тип лестницы {ladder_type}")
            
            # Характеристики вертикальной лестницы
            if ladder.get('height', '').strip():
                characteristics_parts.append(f"высота лестницы {ladder.get('height')} м")
            
            if ladder.get('width', '').strip():
                characteristics_parts.append(f"ширина лестницы {ladder.get('width')} м")
            
            if ladder.get('steps_count', '').strip():
                characteristics_parts.append(f"количество ступеней {ladder.get('steps_count')} (шт.)")
            
            if ladder.get('mount_points', '').strip():
                characteristics_parts.append(f"количество точек крепления {ladder.get('mount_points')} (шт.)")
            
            platform_dims = []
            if ladder.get('platform_length', '').strip():
                platform_dims.append(f"длина {ladder.get('platform_length')} м")
            if ladder.get('platform_width', '').strip():
                platform_dims.append(f"ширина {ladder.get('platform_width')} м")
            if platform_dims:
                characteristics_parts.append(f"размер площадки: {', '.join(platform_dims)}")
            
            if ladder.get('fence_height', '').strip():
                characteristics_parts.append(f"высота ограждений площадки {ladder.get('fence_height')} м")
            
            if ladder.get('wall_distance', '').strip():
                characteristics_parts.append(f"расстояние от стены {ladder.get('wall_distance')} м")
            
            if ladder.get('ground_distance', '').strip():
                characteristics_parts.append(f"расстояние от земли {ladder.get('ground_distance')} м")
            
            if ladder.get('step_distance', '').strip():
                characteristics_parts.append(f"расстояние между ступенями {ladder.get('step_distance')} м")
            
            # Формируем текст: только название + характеристики
            if ladder_name:
                ladder_title = ladder_name
            else:
                ladder_title = f"Лестница №{ladder_num}"
            
            # Выводим параграф с жирным названием
            p = self.document.add_paragraph()
            
            # Название жирным
            run = p.add_run(ladder_title)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Двоеточие и характеристики обычным шрифтом
            if characteristics_parts:
                characteristics_text = ": " + ", ".join(characteristics_parts) + "."
            else:
                characteristics_text = ": характеристики не указаны."
            
            run = p.add_run(characteristics_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            self._format_paragraph(p)
    
    def _add_test_conditions(self, data):
        """Добавляет условия проведения испытаний"""
        heading = self.document.add_heading('Условия проведения испытаний', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Текстовое описание условий без таблицы
        conditions_text = (
            f"Испытания проводились в {data.get('test_time', 'дневное время')} "
            f"при температуре воздуха {data.get('temperature', '')}°C и скорости ветра {data.get('wind_speed', '')} м/с."
        )
        
        p = self.document.add_paragraph(conditions_text)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)
    
    def _add_test_equipment(self, data):
        """Добавляет средства проведения испытаний"""
        heading = self.document.add_heading('Средства проведения испытаний', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Простая строка про оборудование
        equipment_text = "Динамометр ДПУ 0.5-2 заводской №1860 (свидетельство о поверке № С-ВЮМ/23-07-2025/453474803), рулетка измерительная металлическая RGK R-5 заводской № Е5М1270 (свидетельство о поверке № С-ЕВЕ/16-07-2025/448133521)."
        
        p = self.document.add_paragraph(equipment_text)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)
    
    def _add_visual_inspection(self, data):
        """Добавляет визуальный осмотр лестниц (для каждой отдельно)"""
        heading = self.document.add_heading('Визуальный осмотр лестниц', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Получаем список лестниц
        ladders = data.get('ladders', [])
        
        if not ladders:
            # Если нет лестниц - используем старые общие данные
            damage = data.get('damage_found')
            mount = data.get('mount_violation_found')
            weld = data.get('weld_violation_found')
            paint = data.get('paint_compliant', True)
            
            # Создаём параграф
            p = self.document.add_paragraph()
            
            # Внешние повреждения
            damage_text = 'внешние повреждения лестниц обнаружены' if damage else 'внешние повреждения лестниц не обнаружены'
            run = p.add_run(damage_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if damage:
                run.font.bold = True
            
            # Следы нарушения крепления
            mount_text = 'следы нарушения крепления конструкции лестниц к стене здания обнаружены' if mount else 'следы нарушения крепления конструкции лестниц к стене здания не обнаружены'
            run = p.add_run(f", {mount_text}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if mount:
                run.font.bold = True
            
            # Нарушение сварных швов
            weld_text = 'нарушение сварных швов обнаружено' if weld else 'нарушение сварных швов не обнаружено'
            run = p.add_run(f", {weld_text}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if weld:
                run.font.bold = True
            
            # Защитное покрытие
            paint_text = 'защитное покрытие требованиям ГОСТ 9.302 соответствует' if paint else 'защитное покрытие требованиям ГОСТ 9.302 не соответствует'
            run = p.add_run(f", {paint_text}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if not paint:
                run.font.bold = True
            
            run = p.add_run(".")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            self._format_paragraph(p)
        else:
            # Группируем лестницы по одинаковым результатам осмотра
            inspection_groups = {}
            
            for ladder in ladders:
                ladder_num = ladder.get('number', 1)
                ladder_name = ladder.get('name', '')
                
                # Создаём ключ из результатов осмотра
                damage = ladder.get('damage_found', False)
                mount = ladder.get('mount_violation_found', False)
                weld = ladder.get('weld_violation_found', False)
                paint = ladder.get('paint_compliant', True)
                
                key = (damage, mount, weld, paint)
                
                if key not in inspection_groups:
                    inspection_groups[key] = []
                inspection_groups[key].append({'num': ladder_num, 'name': ladder_name})
            
            # Выводим группы
            for (damage, mount, weld, paint), ladder_data_list in inspection_groups.items():
                # Проверяем, всё ли в порядке (нет проблем)
                all_ok = (not damage and not mount and not weld and paint)
                
                # Формируем префикс для группы лестниц
                if all_ok:
                    # Если всё в порядке - не выводим номер/название лестницы
                    ladder_prefix = ""
                else:
                    # Если есть проблемы - выводим с названием
                    if len(ladder_data_list) == 1:
                        ladder_info = ladder_data_list[0]
                        if ladder_info['name']:
                            ladder_prefix = ladder_info['name']
                        else:
                            ladder_prefix = f"Лестница №{ladder_info['num']}"
                    else:
                        # Несколько лестниц с одинаковыми проблемами
                        names = []
                        for info in sorted(ladder_data_list, key=lambda x: x['num']):
                            if info['name']:
                                names.append(info['name'])
                            else:
                                names.append(f"№{info['num']}")
                        ladder_prefix = ', '.join(names)
                
                # Создаём параграф
                p = self.document.add_paragraph()
                
                # Добавляем префикс жирным (если есть)
                if ladder_prefix:
                    run = p.add_run(ladder_prefix)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Добавляем двоеточие обычным шрифтом
                    run = p.add_run(": ")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Внешние повреждения
                damage_text = 'внешние повреждения обнаружены' if damage else 'внешние повреждения не обнаружены'
                run = p.add_run(damage_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                if damage:
                    run.font.bold = True
                
                # Следы нарушения крепления
                mount_text = 'следы нарушения крепления к стене обнаружены' if mount else 'следы нарушения крепления к стене не обнаружены'
                run = p.add_run(f", {mount_text}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                if mount:
                    run.font.bold = True
                
                # Нарушение сварных швов
                weld_text = 'нарушение сварных швов обнаружено' if weld else 'нарушение сварных швов не обнаружено'
                run = p.add_run(f", {weld_text}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                if weld:
                    run.font.bold = True
                
                # Защитное покрытие
                paint_text = 'защитное покрытие требованиям ГОСТ 9.302 соответствует' if paint else 'защитное покрытие требованиям ГОСТ 9.302 не соответствует'
                run = p.add_run(f", {paint_text}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                if not paint:
                    run.font.bold = True
                
                run = p.add_run(".")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                
                self._format_paragraph(p)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)
    
    def _add_load_calculation(self, data):
        """Добавляет раздел расчета величины нагрузки"""
        heading_text = 'Расчет величины нагрузки на лестницу'
        calculation_text = 'Расчет величины нагрузки согласно: ГОСТ Р 53254-2009 «Техника пожарная. Лестницы пожарные наружные стационарные. Ограждения кровли. Общие технические требования. Методы испытаний».'
        
        heading = self.document.add_heading(heading_text, level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        p = self.document.add_paragraph(calculation_text)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)
    
    def _add_test_subjects(self, data):
        """Добавляет раздел 'Испытаниям подлежат'"""
        heading = self.document.add_heading('Испытаниям подлежат', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        test_items = [
            'Балки крепления лестниц к стене (попарно, в месте крепления к лестнице);',
            'Ступени лестницы (в середине ступени) – каждая 5-я ступень;',
            'Ограждения лестницы в точках на расстоянии не более 1,5 м друг от друга по всей высоте лестницы.'
        ]
        
        for item in test_items:
            p = self.document.add_paragraph()
            run = p.add_run('• ' + item)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._format_paragraph(p)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)
    
    def _add_dynamic_table(self, data):
        """Добавляет автоматическую таблицу результатов испытаний"""
        heading = self.document.add_heading('Результаты испытаний', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Определяем максимальную высоту из всех лестниц
        ladders = data.get('ladders', [])
        max_height = 0
        
        if ladders:
            for ladder in ladders:
                try:
                    height = float(str(ladder.get('height', '0')).replace(',', '.'))
                    if height > max_height:
                        max_height = height
                except (ValueError, TypeError):
                    continue
        
        # Если не нашли лестниц, используем старое поле
        if max_height == 0:
            try:
                max_height = float(str(data.get('ladder_height', '0')).replace(',', '.'))
            except (ValueError, TypeError):
                max_height = 0
        
        ladder_height = max_height
        
        # Формируем данные таблицы в зависимости от типа лестницы
        if ladder_height <= 6:
            # П1-1: количество испытываемых точек = 2
            guardrail_points = 2
            guardrail_name = "Ограждения площадки"
        else:
            # П1-2: количество испытываемых точек = 2 + (высота / 1.2)
            guardrail_points = int(2 + (ladder_height / 1.2))
            guardrail_name = "Ограждения лестницы и площадки"
        
        # Данные таблицы
        table_data_auto = [
            ["№ п/п", "Наименование испытываемого элемента", "Кол/во испытываемых точек", "Нагрузка кН (кгс)", "Результаты испытаний"],
            ["1", "Ступени", "3", "1.8 (180)", "Выдержали"],
            ["2", guardrail_name, str(guardrail_points), "0.54 (54)", "Выдержали"],
        ]
        
        # Добавляем балки крепления - берём среднее количество точек крепления
        ladders = data.get('ladders', [])
        total_mount_points = 0
        count = 0
        
        for ladder in ladders:
            mp = ladder.get('mount_points', '')
            if mp:
                try:
                    total_mount_points += int(mp)
                    count += 1
                except ValueError:
                    pass
        
        # Используем старое поле если нет данных в лестницах
        if count == 0:
            mount_points = data.get('mount_points', '')
        else:
            mount_points = str(total_mount_points)
        
        if mount_points:
            try:
                mount_count = int(mount_points)
                # Расчет нагрузки на балки: (Высота * 0.72) / количество_упоров
                load_kn = (ladder_height * 0.72) / mount_count if mount_count > 0 else 0
                load_kgs = load_kn * 100
                table_data_auto.append(["3", "Балки крепления к стене", str(mount_count), f"{load_kn:.2f} ({load_kgs:.0f})", "Выдержали"])
            except (ValueError, ZeroDivisionError):
                table_data_auto.append(["3", "Балки крепления к стене", mount_points, "", "Выдержали"])
        else:
            table_data_auto.append(["3", "Балки крепления к стене", "", "", "Выдержали"])
        
        # Создание таблицы
        rows_count = len(table_data_auto)
        cols_count = 5
        table = self.document.add_table(rows=rows_count, cols=cols_count)
        table.style = 'Table Grid'
        
        # Установка ширины столбцов
        column_widths = {
            0: Inches(1 / 2.54),    # 1 см
            1: Inches(7 / 2.54),    # 7 см
            2: Inches(3 / 2.54),    # 3 см
            3: Inches(3 / 2.54),    # 3 см
            4: Inches(3 / 2.54),    # 3 см
        }
        
        for col_idx, width in column_widths.items():
            for row in table.rows:
                row.cells[col_idx].width = width
        
        # Заполнение данных
        for i, row_data in enumerate(table_data_auto):
            for j, cell_value in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_value)
                
                # Форматирование первой строки (заголовки)
                if i == 0:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Форматирование всех ячеек таблицы
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._format_paragraph(paragraph)
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.color.rgb = RGBColor(0, 0, 0)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)
    
    def _add_conclusions(self, data):
        """Добавляет выводы по результатам испытаний"""
        heading = self.document.add_heading('Выводы по результатам испытаний', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        ladders = data.get('ladders', [])
        ladders_compliance = data.get('ladders_compliance', {})
        project_compliant = data.get('project_compliant', False)
        project_number = data.get('project_number', '')
        
        # Логирование для отладки
        app_logger.info(f"Данные соответствия лестниц: {ladders_compliance}")
        
        # Формирование текста вывода - для каждой лестницы отдельно
        conclusion_parts = []
        
        # Если нет данных по лестницам - используем старую логику
        if not ladders:
            has_damages = (
                data.get('damage_found') or 
                data.get('mount_violation_found') or 
                data.get('weld_violation_found')
            )
            paint_ok = data.get('paint_compliant', True)
            
            if has_damages:
                conclusion_parts.append("Конструкции вертикальных пожарных лестниц не пригодны к эксплуатации.")
            elif not paint_ok:
                conclusion_parts.append("Конструкции вертикальных пожарных лестниц в прочностные испытания выдержали, нагрузка выдерживалась в течение 2 минут. После испытания не имеет трещин, прогибов, изломов. Требуется восстановить защитное покрытие.")
            else:
                conclusion_parts.append("Конструкции вертикальных пожарных лестниц в прочностные испытания выдержали, нагрузка выдерживалась в течение 2 минут. После испытания не имеет трещин, прогибов, изломов. Пригодны к эксплуатации. Соответствует ГОСТ Р 54253-2009 «Техника пожарная. Лестницы пожарные наружные стационарные. Ограждения кровли. Общие технические требования. Методы испытаний».")
        else:
            # Проверяем - все ли лестницы соответствуют требованиям
            all_compliant = True
            
            violation_names = {
                'ladder_width': 'ширина лестницы',
                'step_distance': 'расстояние между ступенями',
                'wall_distance': 'расстояние от стены',
                'ground_distance': 'расстояние от земли',
                'platform_length': 'длина площадки',
                'platform_width': 'ширина площадки',
                'fence_height': 'высота ограждения площадки',
                'ladder_fence': 'ограждение лестницы',
                'mount_distance': 'расстояние между упорами',
                'paint_coating': 'защитное покрытие'
            }
            
            # Собираем информацию о каждой лестнице
            ladder_results = []
            for ladder in ladders:
                ladder_num = ladder.get('number', 1)
                ladder_type = ladder.get('ladder_type', 'vertical')
                if ladder_type not in ('', 'vertical'):
                    app_logger.warning(f"Лестница №{ladder_num}: получен устаревший тип '{ladder_type}', вывод будет сформирован как для вертикальной конструкции")
                
                # Проверяем повреждения этой лестницы
                has_damages = (
                    ladder.get('damage_found') or 
                    ladder.get('mount_violation_found') or 
                    ladder.get('weld_violation_found')
                )
                
                paint_ok = ladder.get('paint_compliant', True)
                
                # Проверяем соответствие ГОСТ для этой лестницы
                compliance_data = ladders_compliance.get(ladder_num, {'compliant': True, 'violations': {}})
                compliant = compliance_data.get('compliant', True)
                gost_violations = compliance_data.get('violations', {})
                
                # Проверяем наличие paint_coating в ГОСТ несоответствиях
                has_paint_violation = not compliant and gost_violations.get('paint_coating', False)
                
                # Если paint_coating в ГОСТ или галочка не стоит - покрытие не ок
                if has_paint_violation or not paint_ok:
                    paint_ok = False
                
                app_logger.info(f"Лестница №{ladder_num}: type={ladder_type}, damages={has_damages}, paint_ok={paint_ok}, compliant={compliant}, has_paint_violation={has_paint_violation}")
                
                ladder_results.append({
                    'has_damages': has_damages,
                    'paint_ok': paint_ok,
                    'compliant': compliant,
                    'gost_violations': gost_violations
                })
            
            ladder_type_name = 'вертикальных пожарных лестниц'
            any_damages = any(l['has_damages'] for l in ladder_results)
            all_paint_ok = all(l['paint_ok'] for l in ladder_results)
            all_compliant_type = all(l['compliant'] for l in ladder_results)
            
            all_violations = set()
            for result in ladder_results:
                for key, value in result['gost_violations'].items():
                    if value:
                        all_violations.add(key)
            
            if any_damages:
                conclusion_parts.append(f"Конструкции {ladder_type_name} не пригодны к эксплуатации.")
            elif not all_compliant_type:
                selected_violations = [violation_names[key] for key in all_violations if key != 'paint_coating']
                base_text = f"Конструкции {ladder_type_name} в прочностные испытания выдержали, нагрузка выдерживалась в течение 2 минут. После испытания не имеет трещин, прогибов, изломов."
                
                if selected_violations:
                    violations_text = ', '.join(selected_violations)
                    conclusion_text = f"{base_text} Требованиям ГОСТ Р 54253-2009 не соответствует ({violations_text})."
                else:
                    if not all_paint_ok:
                        conclusion_text = f"{base_text} Требуется восстановить защитное покрытие."
                    else:
                        conclusion_text = f"{base_text} Требованиям ГОСТ Р 54253-2009 не соответствует."
                
                if not all_paint_ok and selected_violations:
                    conclusion_text = conclusion_text.rstrip('.') + ". Требуется восстановить защитное покрытие."
                
                conclusion_parts.append(conclusion_text)
            elif not all_paint_ok:
                conclusion_parts.append(f"Конструкции {ladder_type_name} в прочностные испытания выдержали, нагрузка выдерживалась в течение 2 минут. После испытания не имеет трещин, прогибов, изломов. Требуется восстановить защитное покрытие.")
            else:
                conclusion_parts.append(f"Конструкции {ladder_type_name} в прочностные испытания выдержали, нагрузка выдерживалась в течение 2 минут. После испытания не имеет трещин, прогибов, изломов. Пригодны к эксплуатации. Соответствует ГОСТ Р 54253-2009 «Техника пожарная. Лестницы пожарные наружные стационарные. Ограждения кровли. Общие технические требования. Методы испытаний».")
        
        conclusion_text = "\n".join(conclusion_parts)
        
        # Соответствие проекту
        if project_compliant and project_number:
            conclusion_text += f"\nСоответствует проекту {project_number}."
        
        app_logger.info(f"Финальный текст вывода: {conclusion_text[:300]}...")
        
        # Добавление текста в документ - разделяем на параграфы по \n
        conclusion_paragraphs = conclusion_text.split('\n')
        
        for i, para_text in enumerate(conclusion_paragraphs):
            if para_text.strip():  # Пропускаем пустые строки
                p = self.document.add_paragraph(para_text.strip())
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                self._format_paragraph(p)
                
                # Добавляем пустую строку между выводами разных лестниц
                if i < len(conclusion_paragraphs) - 1:  # Не после последнего
                    p = self.document.add_paragraph()
                    self._format_paragraph(p)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)
    
    def _add_signatures(self, data):
        """Добавляет раздел с подписями"""
        p = self.document.add_paragraph()
        self._format_paragraph(p)
        p = self.document.add_paragraph()
        self._format_paragraph(p)
        
        # Подпись ответственного лица
        p = self.document.add_paragraph()
        run = p.add_run('Ответственный за проведение испытаний:')
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        p = self.document.add_paragraph()
        run = p.add_run('___________________ / _________________ /')
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        p = self.document.add_paragraph()
        run = p.add_run('        (подпись)                 (Ф.И.О.)')
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
    
    def _generate_filename(self, data):
        """Генерирует имя файла"""
        date_str = data.get('date', datetime.now().strftime('%Y-%m-%d'))
        date_str = date_str.replace('.', '-').replace('/', '-')
        
        # Добавляем время для уникальности
        time_str = datetime.now().strftime('%H-%M-%S')
        
        # Для имени файла используем название первой лестницы или объект
        ladders = data.get('ladders', [])
        if ladders and ladders[0].get('name'):
            base_name = ladders[0]['name']
        else:
            # Берём первые слова из объединённого поля
            base_name = data.get('object_full_address', 'объект')
        
        # Очистка имени от недопустимых символов
        object_name = "".join(c for c in base_name if c.isalnum() or c in (' ', '-', '_')).strip()
        object_name = object_name[:50]  # Ограничение длины
        
        filename = f"Отчёт_{date_str}_{time_str}_{object_name}.docx"
        return filename
    
    def load_contract_data(self, customer_name):
        """
        Загружает данные из договора для указанного заказчика
        
        Args:
            customer_name (str): Имя заказчика
        
        Returns:
            dict: Данные из договора или None
        """
        try:
            # Поиск файла договора
            contract_file = None
            for file in config.CONTRACTS_DIR.glob("*.docx"):
                if customer_name.lower() in file.stem.lower():
                    contract_file = file
                    break
            
            if not contract_file:
                app_logger.warning(f"Договор для заказчика '{customer_name}' не найден")
                return None
            
            # Чтение данных из договора
            doc = Document(str(contract_file))
            
            # Здесь можно реализовать логику извлечения данных
            extracted_data = {
                'object_full_address': '',
                'ladders': [],
            }
            
            app_logger.info(f"Данные из договора '{contract_file.name}' загружены")
            return extracted_data
            
        except Exception as e:
            app_logger.error(f"Ошибка при загрузке договора: {e}")
            return None

