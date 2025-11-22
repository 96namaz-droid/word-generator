"""
Генератор протокола испытания ограждений кровли
"""
from __future__ import annotations

from docx.shared import Pt, RGBColor, Inches

from generators.base_generator import BaseProtocolGenerator


class RoofFenceGenerator(BaseProtocolGenerator):
    """Генерация документа для ограждений кровли"""

    REQUIRED_FIELDS = (
        'date',
        'customer',
        'object_full_address',
        'length',
        'height',
        'mount_points',
    )

    def validate(self) -> None:
        self._require_fields(self.REQUIRED_FIELDS)
        numeric_fields = (
            ('length', 1.0, 500.0),
            ('height', 0.6, 2.5),
            ('mount_points', 2, 500),
        )
        for field, min_value, max_value in numeric_fields:
            value = self._to_float(self.data.get(field))
            if value < min_value or value > max_value:
                raise ValueError(
                    f"Поле '{field}' должно быть в диапазоне {min_value}–{max_value}"
                )

    def generate_doc(self, output_path=None) -> str:
        self._set_document()
        self._add_company_header()
        self._add_title(
            "Протокол испытания ограждений кровли",
            f"от {self.data.get('date', '')}".strip()
        )
        self._add_overview()
        self._add_geometry_section()
        self._add_test_equipment()
        self._add_visual_inspection()
        self._add_load_calculation()
        self._add_load_table()
        self._add_conclusion()
        self._add_signatures()

        filename = self._generate_filename("Протокол_ограждения")
        output = self._resolve_output_path(output_path, filename)
        self.document.save(str(output))
        return str(output)

    # --- Разделы документа -----------------------------------------------------

    def _add_overview(self) -> None:
        self._add_key_value('Заказчик', self.data.get('customer', ''))
        self._add_key_value(
            'Адрес/наименование объекта',
            self.data.get('object_full_address', '')
        )
        self.document.add_paragraph()

    def _add_geometry_section(self) -> None:
        """Добавляет характеристики испытываемых конструкций"""
        # Собираем данные
        fence_name = self.data.get('fence_name', '').strip()
        length = self.data.get('length', '').strip()
        height = self.data.get('height', '').strip()
        mount_points = self.data.get('mount_points', '').strip()
        parapet_height = self.data.get('parapet_height', '').strip()

        # Формируем текст характеристики
        
        # 1. Название
        if fence_name:
            name_part = f"Ограждения кровли {fence_name}"
        else:
            name_part = "Ограждения кровли"
        
        # 2. Конструкция (сразу после названия через пробел)
        construction_part = "Конструкция ограждений – металлическая конструкция с вертикальными и горизонтальными ограждающими элементами"
        
        # 3. Длина (через тире)
        length_part = f"- длиной {length} м.п." if length else ""
        
        # 4-6. Остальные характеристики через запятую
        other_parts = []
        if parapet_height:
            other_parts.append(f"высота от плоскости кровли {parapet_height} м")
        if height:
            other_parts.append(f"высота ограждений {height} м")
        if mount_points:
            other_parts.append(f"количество точек креплений {mount_points} шт")
        
        # Собираем текст: название + конструкция через пробел
        main_part = f"{name_part} {construction_part}"
        
        # Собираем все части (длина отдельно с тире, остальное через запятые)
        all_parts = []
        if length_part:
            all_parts.append(length_part)
        all_parts.extend(other_parts)
        
        # Формируем полный текст с заголовком
        if all_parts:
            # Между main_part и all_parts ставим пробел только если длина есть (она уже содержит тире)
            # Если длины нет, просто join через запятые
            content_text = f"{main_part} {', '.join(all_parts)}."
        else:
            content_text = f"{main_part}."
        
        # Создаём абзац с жирным заголовком
        paragraph = self.document.add_paragraph()
        
        # Добавляем жирный заголовок
        bold_run = paragraph.add_run("Характеристики испытываемых конструкций: ")
        bold_run.font.name = 'Times New Roman'
        bold_run.font.size = Pt(11)
        bold_run.font.bold = True
        bold_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Добавляем обычный текст с характеристиками
        normal_run = paragraph.add_run(content_text)
        normal_run.font.name = 'Times New Roman'
        normal_run.font.size = Pt(11)
        normal_run.font.bold = False
        normal_run.font.color.rgb = RGBColor(0, 0, 0)
        
        self._format_paragraph(paragraph)
        self.document.add_paragraph()

    def _add_test_equipment(self) -> None:
        """Добавляет средства проведения испытаний"""
        heading = self.document.add_heading('Средства проведения испытаний', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Текст про оборудование
        equipment_text = (
            "Динамометр ДПУ 0.5-2 заводской №1860 (свидетельство о поверке № С-ВЮМ/23-07-2025/453474803), "
            "рулетка измерительная металлическая RGK R-5 заводской № Е5М1270 (свидетельство о поверке № С-ЕВЕ/16-07-2025/448133521)."
        )
        
        p = self.document.add_paragraph(equipment_text)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)

    def _add_visual_inspection(self) -> None:
        """Добавляет визуальный осмотр ограждений"""
        heading = self.document.add_heading('Визуальный осмотр ограждений', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Получаем данные о визуальном осмотре
        damage = self.data.get('damage_found', False)
        mount = self.data.get('mount_violation_found', False)
        weld = self.data.get('weld_violation_found', False)
        paint = self.data.get('paint_compliant', True)
        
        # Создаём параграф
        p = self.document.add_paragraph()
        
        # Внешние повреждения
        damage_text = 'внешние повреждения не обнаружены' if not damage else 'внешние повреждения обнаружены'
        run = p.add_run(damage_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if damage:
            run.font.bold = True
        
        # Следы нарушения крепления (для ограждений - "к стене")
        mount_text = 'следы нарушения крепления к стене не обнаружены' if not mount else 'следы нарушения крепления к стене обнаружены'
        run = p.add_run(f", {mount_text}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if mount:
            run.font.bold = True
        
        # Нарушение сварных швов
        weld_text = 'нарушение сварных швов не обнаружено' if not weld else 'нарушение сварных швов обнаружено'
        run = p.add_run(f", {weld_text}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if weld:
            run.font.bold = True
        
        # Защитное покрытие
        paint_text = 'защитное покрытие требованиям ГОСТ 9.302 соответствует' if paint else 'защитное покрытие требованиям ГОСТ 9.302 не соответствует'
        run = p.add_run(f", {paint_text}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if not paint:
            run.font.bold = True
        
        run = p.add_run(".")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        self._format_paragraph(p)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)

    def _add_load_calculation(self) -> None:
        """Добавляет раздел расчета величины нагрузки"""
        heading_text = 'Расчет величины нагрузки на ограждения'
        calculation_text = (
            'Расчет величины нагрузки согласно: ГОСТ Р 53254-2009 «Техника пожарная. '
            'Лестницы пожарные наружные стационарные. Ограждения кровли. '
            'Общие технические требования. Методы испытаний».'
        )
        
        heading = self.document.add_heading(heading_text, level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        p = self.document.add_paragraph(calculation_text)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
        self._format_paragraph(p)
        
        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)

    def _add_load_table(self) -> None:
        """Добавляет таблицу результатов нагрузочных испытаний"""
        heading = self.document.add_heading('Результаты испытаний', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        length = self._to_float(self.data.get('length'))
        fence_name = self.data.get('fence_name', '').strip()
        
        # Название из блока характеристик
        if fence_name:
            element_name = f"Ограждения кровли {fence_name}"
        else:
            element_name = "Ограждения кровли"
        
        # Количество испытываемых точек = (длина в погонных метрах / 10) + 3
        test_points = int((length / 10) + 3)
        
        # Данные таблицы - только одна строка
        table_data = [
            ["№ п/п", "Наименование испытываемого элемента", "Кол/во испытываемых точек", "Нагрузка кН (кгс)", "Результаты испытаний"],
            ["1", element_name, str(test_points), "0.54 (54)", "Выдержали"],
        ]

        # Создание таблицы
        rows_count = len(table_data)
        cols_count = 5
        table = self.document.add_table(rows=rows_count, cols=cols_count)
        table.style = 'Table Grid'

        # Установка ширины столбцов (как в vertical_ladder.py)
        column_widths = {
            0: Inches(1 / 2.54),    # 1 см
            1: Inches(7 / 2.54),    # 7 см
            2: Inches(3 / 2.54),    # 3 см
            3: Inches(3 / 2.54),    # 3 см
            4: Inches(3 / 2.54),    # 3 см
        }

        for col_idx, width in column_widths.items():
            for row in table.rows:
                row.cells[col_idx].width = width

        # Заполнение данных
        for i, row_data in enumerate(table_data):
            for j, cell_value in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = str(cell_value)

                # Форматирование первой строки (заголовки)
                if i == 0:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                    cell.paragraphs[0].runs[0].font.size = Pt(11)
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)

        # Форматирование всех ячеек таблицы
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._format_paragraph(paragraph)
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)

        p = self.document.add_paragraph()  # Пустая строка
        self._format_paragraph(p)

    def _add_conclusion(self) -> None:
        heading = self.document.add_paragraph()
        heading_run = heading.add_run('Выводы по результатам испытаний: ')
        heading_run.font.name = 'Times New Roman'
        heading_run.font.size = Pt(11)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Получаем данные о визуальном осмотре
        damage_found = self.data.get('damage_found', False)
        mount_violation_found = self.data.get('mount_violation_found', False)
        weld_violation_found = self.data.get('weld_violation_found', False)
        paint_compliant = self.data.get('paint_compliant', True)
        
        # Словарь для названий проблем
        problems = []
        problem_names = {
            'damage': 'внешние повреждения',
            'mount': 'следы нарушения крепления',
            'weld': 'нарушения сварных швов'
        }
        
        if damage_found:
            problems.append(problem_names['damage'])
        if mount_violation_found:
            problems.append(problem_names['mount'])
        if weld_violation_found:
            problems.append(problem_names['weld'])
        
        # Базовый вывод (выводится всегда, если нет проблем)
        paragraph = self.document.add_paragraph()
        
        # Если есть проблемы - выводим что не пригодны
        if problems:
            problem_text = ', '.join(problems)
            conclusion_text = f"Ограждения кровли не пригодны к эксплуатации ({problem_text})."
            
            conclusion_run = paragraph.add_run(conclusion_text)
            conclusion_run.font.name = 'Times New Roman'
            conclusion_run.font.size = Pt(11)
            conclusion_run.font.bold = False
            conclusion_run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            # Базовый вывод если нет проблем
            base_text = (
                "Конструкции ограждений кровли прочностные испытания выдержали, "
                "нагрузка выдерживалась в течении 2-х минут. После испытания не имеет трещин, "
                "прогибов, изломов. Соответствует требованиям ГОСТ Р 53254-2009 "
                "«Техника пожарная. Лестницы пожарные наружные стационарные. Ограждения кровли. "
                "Общие технические требования. Методы испытаний». Пригодны к эксплуатации."
            )
            
            conclusion_run = paragraph.add_run(base_text)
            conclusion_run.font.name = 'Times New Roman'
            conclusion_run.font.size = Pt(11)
            conclusion_run.font.bold = False
            conclusion_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Если защитное покрытие не соответствует - добавляем требование
            if not paint_compliant:
                paint_paragraph = self.document.add_paragraph()
                paint_run = paint_paragraph.add_run("Требуется восстановить защитное покрытие.")
                paint_run.font.name = 'Times New Roman'
                paint_run.font.size = Pt(11)
                paint_run.font.bold = False
                paint_run.font.color.rgb = RGBColor(0, 0, 0)
        
        self.document.add_paragraph()

    # --- Helpers ---------------------------------------------------------------

    @staticmethod
    def _to_float(value, default=0.0) -> float:
        try:
            return float(str(value).replace(',', '.'))
        except (TypeError, ValueError, AttributeError):
            return default


