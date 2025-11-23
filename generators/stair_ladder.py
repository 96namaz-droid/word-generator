"""
Генератор протокола испытания маршевых пожарных лестниц
"""
from __future__ import annotations
import math

from docx.shared import Pt, RGBColor, Mm

from generators.base_generator import BaseProtocolGenerator


class StairLadderGenerator(BaseProtocolGenerator):
    """Генерация документа для маршевых лестниц"""

    REQUIRED_FIELDS = (
        'date',
        'customer',
        'object_full_address',
        'march_width',
        'march_length',
        'step_width',
        'step_distance',
        'steps_count',
        'platform_fence_height',
        'platform_length',
        'platform_width',
        'platform_fence_height_2',
    )

    def validate(self) -> None:
        # Проверяем, есть ли марши в новом формате
        marches = self.data.get('marches', [])
        
        if marches:
            # Новый формат - валидируем каждый марш
            numeric_fields = (
                ('march_width', 0.5, 10.0),
                ('march_length', 0.5, 50.0),
                ('step_width', 0.15, 1.0),
                ('step_distance', 0.15, 0.5),
                ('steps_count', 1, 100),
                ('march_fence_height', 0.5, 2.5),
                ('platform_fence_height', 0.5, 2.5),
                ('platform_length', 0.5, 10.0),
                ('platform_width', 0.5, 10.0),
                ('platform_ground_distance', 0.0, 50.0),  # Опциональное поле
            )
            
            # Базовые обязательные поля для марша и площадки
            march_required_fields = [
                'march_width', 'march_length', 'step_width', 'step_distance', 
                'steps_count', 'march_fence_height'
            ]
            platform_required_fields = [
                'platform_fence_height', 'platform_length', 'platform_width'
            ]
            
            for march in marches:
                march_num = march.get('number', '?')
                has_march = march.get('has_march', True)
                has_platform = march.get('has_platform', True)
                
                # Проверяем обязательные поля марша, если марш присутствует
                if has_march:
                    missing_march = [f for f in march_required_fields if not str(march.get(f, '')).strip()]
                    if missing_march:
                        raise ValueError(
                            f"Марш №{march_num}: Не заполнены обязательные поля марша: {', '.join(missing_march)}"
                        )
                
                # Проверяем обязательные поля площадки, если площадка присутствует
                if has_platform:
                    missing_platform = [f for f in platform_required_fields if not str(march.get(f, '')).strip()]
                    if missing_platform:
                        raise ValueError(
                            f"Площадка №{march_num}: Не заполнены обязательные поля площадки: {', '.join(missing_platform)}"
                        )
                
                # Проверяем, что хотя бы марш или площадка присутствуют
                if not has_march and not has_platform:
                    raise ValueError(
                        f"Элемент №{march_num}: Должен быть заполнен хотя бы марш или площадка"
                    )
                
                # Проверяем числовые поля марша, если марш присутствует
                if has_march:
                    march_numeric_fields = [
                        ('march_width', 0.5, 10.0),
                        ('march_length', 0.5, 50.0),
                        ('step_width', 0.15, 1.0),
                        ('step_distance', 0.15, 0.5),
                        ('steps_count', 1, 100),
                        ('march_fence_height', 0.5, 2.5),
                    ]
                    for field, min_value, max_value in march_numeric_fields:
                        value = self._to_float(march.get(field))
                        if value > 0 and (value < min_value or value > max_value):
                            raise ValueError(
                                f"Марш №{march_num}, поле '{field}': должно быть в диапазоне {min_value}–{max_value}"
                            )
                
                # Проверяем числовые поля площадки, если площадка присутствует
                if has_platform:
                    platform_numeric_fields = [
                        ('platform_fence_height', 0.5, 2.5),
                        ('platform_length', 0.5, 10.0),
                        ('platform_width', 0.5, 10.0),
                        ('platform_ground_distance', 0.0, 50.0),
                    ]
                    for field, min_value, max_value in platform_numeric_fields:
                        value = self._to_float(march.get(field))
                        if value > 0 and (value < min_value or value > max_value):
                            raise ValueError(
                                f"Площадка №{march_num}, поле '{field}': должно быть в диапазоне {min_value}–{max_value}"
                            )
            
            # Проверяем базовые обязательные поля
            self._require_fields(['date', 'customer', 'object_full_address'])
        else:
            # Старый формат - валидация как раньше
            self._require_fields(self.REQUIRED_FIELDS)
            numeric_fields = (
                ('march_width', 0.5, 10.0),
                ('march_length', 0.5, 50.0),
                ('step_width', 0.15, 1.0),
                ('step_distance', 0.15, 0.5),
                ('steps_count', 1, 100),
                ('platform_fence_height', 0.5, 2.5),
                ('platform_length', 0.5, 10.0),
                ('platform_width', 0.5, 10.0),
                ('platform_fence_height_2', 0.5, 2.5),
                ('platform_ground_distance', 0.0, 50.0),  # Опциональное поле
            )
            for field, min_value, max_value in numeric_fields:
                value = self._to_float(self.data.get(field))
                # Проверяем только если поле заполнено (не пустое)
                if value > 0 and (value < min_value or value > max_value):
                    raise ValueError(
                        f"Поле '{field}' должно быть в диапазоне {min_value}–{max_value}"
                    )

    def generate_doc(self, output_path=None) -> str:
        self._set_document()
        self._add_company_header()
        self._add_title(
            "Протокол испытания маршевых лестниц",
            f"от {self.data.get('date', '')}".strip()
        )
        self._add_overview()
        self._add_parameters_section()
        self._add_environment_section()
        self._add_test_equipment_section()
        self._add_visual_inspection_section()
        self._add_requirements_section()
        self._add_load_table()
        self._add_conclusion()
        self._add_signatures()

        filename = self._generate_filename("Protocol_stair")
        output = self._resolve_output_path(output_path, filename)
        self.document.save(str(output))
        return str(output)

    # --- Разделы документа -----------------------------------------------------

    def _add_overview(self) -> None:
        self._add_key_value('Заказчик', self.data.get('customer', ''))
        self._add_key_value(
            'Адрес/наименование объекта',
            self.data.get('object_full_address', '')
        )
        additional = self.data.get('object_description', '').strip()
        if additional:
            self._add_key_value('Описание площадки', additional)
        self.document.add_paragraph()

    def _add_parameters_section(self) -> None:
        paragraph = self.document.add_heading('Характеристика испытываемого объекта', level=1)
        self._format_paragraph(paragraph)
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Получаем название лестницы или используем дефолт
        ladder_name = self.data.get('ladder_name', '').strip()
        if not ladder_name:
            ladder_name = "Маршевая лестница №1"

        # Выводим название лестницы и тип
        name_text = f"{ladder_name}. Тип П-2."
        name_paragraph = self.document.add_paragraph(name_text)
        self._format_paragraph(name_paragraph)

        # Получаем список маршей
        marches = self.data.get('marches', [])
        
        # Если есть марши в новом формате, выводим их
        if marches:
            elements_parts = []
            march_fence_heights = set()  # Для сбора высот ограждений маршей
            platform_fence_heights = set()  # Для сбора высот ограждений площадок
            
            for march in marches:
                march_num = march.get('number', '?')
                has_march = march.get('has_march', True)
                has_platform = march.get('has_platform', True)
                
                element_parts = []
                
                # Добавляем параметры марша, если марш присутствует
                if has_march:
                    march_width = march.get('march_width', '').strip()
                    march_length = march.get('march_length', '').strip()
                    step_width = march.get('step_width', '').strip()
                    step_distance = march.get('step_distance', '').strip()
                    steps_count = march.get('steps_count', '').strip()
                    march_fence_height = march.get('march_fence_height', '').strip()
                    
                    march_params = []
                    if march_width:
                        march_params.append(f"ширина марша №{march_num} – {march_width}м")
                    if steps_count:
                        march_params.append(f"кол-во ступеней марш №{march_num} - {steps_count}шт.")
                    # Дополнительные параметры марша (если нужны)
                    if march_length:
                        march_params.append(f"длина марша №{march_num} – {march_length}м")
                    if step_width:
                        march_params.append(f"ширина ступени марша №{march_num} – {step_width}м")
                    if step_distance:
                        march_params.append(f"расстояние между ступенями марша №{march_num} – {step_distance}м")
                    
                    if march_params:
                        element_parts.extend(march_params)
                    
                    if march_fence_height:
                        march_fence_heights.add(march_fence_height)
                
                # Добавляем параметры площадки, если площадка присутствует
                if has_platform:
                    platform_length = march.get('platform_length', '').strip()
                    platform_width = march.get('platform_width', '').strip()
                    platform_fence_height = march.get('platform_fence_height', '').strip()
                    
                    if platform_length and platform_width:
                        element_parts.append(f"площадка №{march_num} размером – {platform_length}*{platform_width}м.")
                    
                    if platform_fence_height:
                        platform_fence_heights.add(platform_fence_height)
                
                if element_parts:
                    elements_parts.append(", ".join(element_parts))
            
            # Формируем текст элементов
            if elements_parts:
                # Каждый элемент заканчивается точкой, соединяем через "; "
                # Для всех кроме последнего заменяем точку на точку с запятой
                parts_list = elements_parts.copy()
                if len(parts_list) > 1:
                    # Все кроме последнего должны заканчиваться точкой с запятой
                    for i in range(len(parts_list) - 1):
                        if parts_list[i].endswith('.'):
                            parts_list[i] = parts_list[i][:-1] + ";"
                elements_text = " ".join(parts_list)
                
                # Добавляем информацию о высоте ограждений в ту же строку
                fence_text_parts = []
                march_fence = sorted(march_fence_heights)[0] if march_fence_heights else None
                platform_fence = sorted(platform_fence_heights)[0] if platform_fence_heights else None
                
                if march_fence and platform_fence:
                    # Если есть оба значения
                    if march_fence == platform_fence:
                        # Одинаковые - объединенный формат
                        fence_text_parts.append(f"высота ограждений марша и площадки лестницы {march_fence}м")
                    else:
                        # Разные - отдельно
                        fence_text_parts.append(f"высота ограждений марша {march_fence}м")
                        fence_text_parts.append(f"высота ограждений площадки лестницы {platform_fence}м")
                elif march_fence:
                    # Только марш
                    fence_text_parts.append(f"высота ограждений марша {march_fence}м")
                elif platform_fence:
                    # Только площадка
                    fence_text_parts.append(f"высота ограждений площадки лестницы {platform_fence}м")
                
                if fence_text_parts:
                    # Добавляем высоту ограждений в конец строки
                    fence_text = ", ".join(fence_text_parts)
                    elements_text = elements_text.rstrip('.') + ", " + fence_text
                else:
                    # Убираем точку в конце, если нет высоты ограждений
                    elements_text = elements_text.rstrip('.')
                
                # Добавляем количество точек крепления в ту же строку
                mount_points = self.data.get('mount_points', '').strip()
                if mount_points:
                    elements_text += f", количество точек крепления {mount_points} шт."
                
                # Убеждаемся, что текст заканчивается точкой
                if not elements_text.endswith('.'):
                    elements_text += "."
                
                elements_paragraph = self.document.add_paragraph(f"Элементы лестницы: {elements_text}")
                self._format_paragraph(elements_paragraph)
        else:
            # Старый формат - один марш (для обратной совместимости)
            march_width = self.data.get('march_width', '').strip()
            steps_count = self.data.get('steps_count', '').strip()
            platform_length = self.data.get('platform_length', '').strip()
            platform_width = self.data.get('platform_width', '').strip()
            march_fence_height = self.data.get('march_fence_height', '').strip()
            platform_fence_height = self.data.get('platform_fence_height', '').strip()
            if not platform_fence_height:
                platform_fence_height = self.data.get('platform_fence_height_2', '').strip()
            
            elements_parts = []
            if march_width and steps_count:
                elements_parts.append(f"Ширина марша №1 – {march_width}м")
                elements_parts.append(f"кол-во ступеней марш №1 - {steps_count}шт.")
            
            if platform_length and platform_width:
                elements_parts.append(f"площадка №1 размером – {platform_length}*{platform_width}м")
            
            if elements_parts:
                elements_text = ", ".join(elements_parts)
                
                # Добавляем информацию о высоте ограждений в ту же строку
                fence_text_parts = []
                
                if march_fence_height and platform_fence_height:
                    # Если есть оба значения
                    if march_fence_height == platform_fence_height:
                        # Одинаковые - объединенный формат
                        fence_text_parts.append(f"высота ограждений марша и площадки лестницы {march_fence_height}м")
                    else:
                        # Разные - отдельно
                        fence_text_parts.append(f"высота ограждений марша {march_fence_height}м")
                        fence_text_parts.append(f"высота ограждений площадки лестницы {platform_fence_height}м")
                elif march_fence_height:
                    # Только марш
                    fence_text_parts.append(f"высота ограждений марша {march_fence_height}м")
                elif platform_fence_height:
                    # Только площадка
                    fence_text_parts.append(f"высота ограждений площадки лестницы {platform_fence_height}м")
                
                if fence_text_parts:
                    # Добавляем высоту ограждений в конец строки
                    fence_text = ", ".join(fence_text_parts)
                    elements_text = elements_text.rstrip('.') + ", " + fence_text
                else:
                    # Убираем точку в конце, если нет высоты ограждений
                    elements_text = elements_text.rstrip('.')
                
                # Добавляем количество точек крепления в ту же строку
                mount_points = self.data.get('mount_points', '').strip()
                if mount_points:
                    elements_text += f", количество точек крепления {mount_points} шт."
                
                # Убеждаемся, что текст заканчивается точкой
                if not elements_text.endswith('.'):
                    elements_text += "."
                
                elements_paragraph = self.document.add_paragraph(f"Элементы лестницы: {elements_text}")
                self._format_paragraph(elements_paragraph)
        
        self.document.add_paragraph()

    def _add_environment_section(self) -> None:
        paragraph = self.document.add_heading('Условия проведения испытаний', level=1)
        self._format_paragraph(paragraph)
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

        test_time = self.data.get('test_time', 'дневное время')
        temperature = self.data.get('temperature', '')
        wind_speed = self.data.get('wind_speed', '')

        text = (
            f"Испытания проводились в {test_time}, "
            f"температура воздуха {temperature} °C, "
            f"скорость ветра {wind_speed} м/с."
        )
        paragraph = self.document.add_paragraph(text)
        self._format_paragraph(paragraph)
        self.document.add_paragraph()

    def _add_test_equipment_section(self) -> None:
        """Добавляет блок средств испытания"""
        heading = self.document.add_heading('Средства испытания', level=1)
        self._format_paragraph(heading)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

        equipment_text = (
            "Динамометр ДПУ 0.5-2 заводской №1860 (свидетельство о поверке № С-ВЮМ/23-07-2025/453474803), "
            "рулетка измерительная металлическая RGK R-5 заводской № Е5М0170 (свидетельство о поверке № С-ЕВЕ/16-07-2025/448133521)."
        )

        paragraph = self.document.add_paragraph(equipment_text)
        self._format_paragraph(paragraph)
        self.document.add_paragraph()

    def _add_visual_inspection_section(self) -> None:
        """Добавляет блок визуального осмотра лестниц"""
        heading = self.document.add_heading('Визуальный осмотр лестниц', level=1)
        self._format_paragraph(heading)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Получаем данные о визуальном осмотре
        damage = self.data.get('damage_found', False)
        mount = self.data.get('mount_violation_found', False)
        weld = self.data.get('weld_violation_found', False)
        paint = self.data.get('paint_compliant', True)

        paragraph = self.document.add_paragraph()

        # Внешние повреждения
        damage_text = 'внешние повреждения конструкций лестницы обнаружены' if damage else 'внешние повреждения конструкций лестницы не обнаружены'
        run = paragraph.add_run(damage_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if damage:
            run.font.bold = True

        # Следы нарушения крепления
        mount_text = 'следы нарушения крепления конструкции лестницы к стене здания обнаружены' if mount else 'следы нарушения крепления конструкции лестницы к стене здания не обнаружены'
        run = paragraph.add_run(f", {mount_text}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if mount:
            run.font.bold = True

        # Нарушение сварных швов
        weld_text = 'нарушение сварных швов обнаружено' if weld else 'нарушение сварных швов не обнаружено'
        run = paragraph.add_run(f", {weld_text}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if weld:
            run.font.bold = True

        # Защитное покрытие
        paint_text = 'защитное покрытие требованиям ГОСТ 9.302 соответствует' if paint else 'защитное покрытие требованиям ГОСТ 9.302 не соответствует'
        run = paragraph.add_run(f", {paint_text}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        if not paint:
            run.font.bold = True

        run = paragraph.add_run(".")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)

        self._format_paragraph(paragraph)
        self.document.add_paragraph()

    def _add_requirements_section(self) -> None:
        heading = self.document.add_heading('Расчет величины нагрузки на лестницу', level=1)
        self._format_paragraph(heading)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

        text = "Расчет величины нагрузки на лестницу согласно: ГОСТ Р53254-2009г."
        paragraph = self.document.add_paragraph(text)
        self._format_paragraph(paragraph)
        self.document.add_paragraph()

    def _add_load_table(self) -> None:
        heading = self.document.add_heading('Результаты испытаний', level=1)
        self._format_paragraph(heading)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Получаем список маршей
        marches = self.data.get('marches', [])
        
        # Список для хранения данных маршей (для формирования отдельных строк в таблице)
        march_loads_data = []  # Список кортежей (номер марша, нагрузка, количество точек)
        # Список для хранения данных площадок (для формирования отдельных строк в таблице)
        platform_loads_data = []  # Список кортежей (номер площадки, нагрузка, количество точек)
        
        if marches:
            # Новый формат - несколько маршей
            total_step_points = 0
            total_rail_points = 0
            total_platform_points = 0
            march_count = 0
            platform_count = 0
            mount_points = self._to_float(self.data.get('mount_points', 0))
            
            for march in marches:
                has_march = march.get('has_march', True)
                has_platform = march.get('has_platform', True)
                
                # Расчет для марша, если он присутствует
                if has_march:
                    steps_count = max(1, int(self._to_float(march.get('steps_count', 0))))
                    # Количество точек для ступеней: (количество ступеней / 5) + 1, минимум 2
                    step_points = max(2, int(steps_count / 5) + 1)
                    total_step_points += step_points
                    march_count += 1
                    
                    # Расчет нагрузки для каждого марша отдельно
                    march_num = march.get('number', march_count)
                    march_length = self._to_float(march.get('march_length', 0))
                    # Высота от земли берется из соответствующей площадки
                    ground_height = self._to_float(march.get('platform_ground_distance', 0))
                    
                    # Расчет нагрузки для марша №N
                    # Формула: длина марша * 1.2 * 1.5 * sqrt(длина марша^2 - высота от земли^2) / (0.5 * количество точек крепления)
                    if mount_points > 0 and march_length > 0:
                        # Вычисляем подкоренное выражение
                        sqrt_value = march_length * march_length - ground_height * ground_height
                        if sqrt_value > 0:
                            sqrt_result = math.sqrt(sqrt_value)
                            # Расчет нагрузки
                            march_load = (march_length * 1.2 * 1.5 * sqrt_result) / (0.5 * mount_points)
                            march_load = round(march_load, 2)
                        else:
                            march_load = 1.5  # Дефолтное значение, если подкоренное выражение <= 0
                    else:
                        march_load = 1.5  # Дефолтное значение
                    
                    # Сохраняем данные марша для таблицы (количество точек = 2 для каждого марша)
                    march_loads_data.append((march_num, march_load, 2))
                
                # Расчет для площадки, если она присутствует
                if has_platform:
                    platform_count += 1
                    platform_num = march.get('number', platform_count)
                    
                    # Расчет нагрузки для каждой площадки отдельно
                    platform_length = self._to_float(march.get('platform_length', 0))
                    platform_width = self._to_float(march.get('platform_width', 0))
                    
                    # Расчет нагрузки для площадки №N
                    # Формула: длина площадки * ширина площадки * 1.2 * 1.5 / (0.5 * количество точек крепления)
                    if mount_points > 0 and platform_length > 0 and platform_width > 0:
                        # Расчет нагрузки
                        platform_load = (platform_length * platform_width * 1.2 * 1.5) / (0.5 * mount_points)
                        platform_load = round(platform_load, 2)
                    else:
                        platform_load = 2.0  # Дефолтное значение
                    
                    # Сохраняем данные площадки для таблицы (количество точек = 1 для каждой площадки)
                    platform_loads_data.append((platform_num, platform_load, 1))
            
            # Ограждения маршей: минимум 6 точек на каждый марш
            rail_points = max(6, 6 * march_count)
            
            # Ограждения площадок: 4 точки на каждую площадку
            platform_points = 4 * platform_count if platform_count > 0 else 0
            
            step_points = total_step_points
        else:
            # Старый формат - один марш (для обратной совместимости)
            steps_count = max(1, int(self._to_float(self.data.get('steps_count', 0))))
            
            # Количество точек для ступеней: (количество ступеней / 5) + 1, минимум 2
            step_points = max(2, int(steps_count / 5) + 1)
            
            # Ограждения маршей: минимум 6 точек
            rail_points = 6
            
            # Ограждения площадок: 4 точки
            platform_points = 4
            
            # Расчет для "Марш лестницы": количество точек = 2
            march_length = self._to_float(self.data.get('march_length', 0))
            ground_height = self._to_float(self.data.get('platform_ground_distance', 0))
            mount_points = self._to_float(self.data.get('mount_points', 0))
            
            # Расчет нагрузки для одного марша (старый формат)
            if mount_points > 0 and march_length > 0:
                # Вычисляем подкоренное выражение
                sqrt_value = march_length * march_length - ground_height * ground_height
                if sqrt_value > 0:
                    sqrt_result = math.sqrt(sqrt_value)
                    # Расчет нагрузки
                    march_load = (march_length * 1.2 * 1.5 * sqrt_result) / (0.5 * mount_points)
                    march_load = round(march_load, 2)
                else:
                    march_load = 1.5  # Дефолтное значение
            else:
                march_load = 1.5  # Дефолтное значение
            
            # Сохраняем данные марша для таблицы (количество точек = 2)
            march_loads_data.append((1, march_load, 2))
            
            # Расчет для "Площадка лестницы": количество точек = 1
            platform_length = self._to_float(self.data.get('platform_length', 0))
            platform_width = self._to_float(self.data.get('platform_width', 0))
            
            if mount_points > 0 and platform_length > 0 and platform_width > 0:
                # Расчет нагрузки
                platform_load = (platform_length * platform_width * 1.2 * 1.5) / (0.5 * mount_points)
                platform_load = round(platform_load, 2)
            else:
                platform_load = 2.0  # Дефолтное значение
            
            # Сохраняем данные площадки для таблицы (количество точек = 1)
            platform_loads_data.append((1, platform_load, 1))

        # Заголовки таблицы согласно рисунку
        headers = (
            "№ п/п",
            "Наименование испытываемого элемента",
            "Кол/во испытываемых точек",
            "Нагрузка кН (кгс)",
            "Результаты испытаний"
        )

        # Формируем базовые строки таблицы
        rows = [
            ("1", "Ступени маршевых лестниц", str(step_points), "1.8 кН (180 кгс)", "Выдержали"),
            ("2", "Ограждения площадок", str(platform_points), "0.54 кН (54 кгс)", "Выдержали"),
            ("3", "Ограждения маршей", str(rail_points), "0.54 кН (54 кгс)", "Выдержали"),
        ]
        
        # Добавляем строки для каждого марша
        row_num = 4
        for march_num, march_load, march_points in march_loads_data:
            march_load_kg = int(march_load * 100)  # Переводим кН в кгс
            if len(march_loads_data) > 1:
                # Если несколько маршей, указываем номер
                rows.append((str(row_num), f"Марш лестницы №{march_num}", str(march_points), f"{march_load:.2f} кН ({march_load_kg} кгс)", "Выдержали"))
            else:
                # Если один марш, без номера
                rows.append((str(row_num), "Марш лестницы", str(march_points), f"{march_load:.2f} кН ({march_load_kg} кгс)", "Выдержали"))
            row_num += 1
        
        # Добавляем строки для каждой площадки
        for platform_num, platform_load, platform_points in platform_loads_data:
            platform_load_kg = int(platform_load * 100)  # Переводим кН в кгс
            if len(platform_loads_data) > 1:
                # Если несколько площадок, указываем номер
                rows.append((str(row_num), f"Площадка лестницы №{platform_num}", str(platform_points), f"{platform_load:.2f} кН ({platform_load_kg} кгс)", "Выдержали"))
            else:
                # Если одна площадка, без номера
                rows.append((str(row_num), "Площадка лестницы", str(platform_points), f"{platform_load:.2f} кН ({platform_load_kg} кгс)", "Выдержали"))
            row_num += 1

        # Создаем таблицу с заданными ширинами столбцов
        if not self.document:
            return

        table = self.document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False

        # Заполняем заголовки
        for idx, header in enumerate(headers):
            cell = table.cell(0, idx)
            cell.text = header
            for paragraph in cell.paragraphs:
                self._format_paragraph(paragraph)
                if paragraph.runs:
                    paragraph.runs[0].font.bold = True
                    paragraph.runs[0].font.name = 'Times New Roman'
                    paragraph.runs[0].font.size = Pt(10)

        # Заполняем данные
        for row_idx, row_data in enumerate(rows, start=1):
            for col_idx, value in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    self._format_paragraph(paragraph)
                    if paragraph.runs:
                        paragraph.runs[0].font.name = 'Times New Roman'
                        paragraph.runs[0].font.size = Pt(10)

        # Устанавливаем ширины столбцов после заполнения: 10мм, 70мм, 30мм, 30мм, 30мм
        # В python-docx нужно устанавливать ширину для каждой ячейки в столбце
        column_widths = [Mm(10), Mm(70), Mm(30), Mm(30), Mm(30)]
        for col_idx, width in enumerate(column_widths):
            for row in table.rows:
                row.cells[col_idx].width = width

        self.document.add_paragraph()

    def _add_conclusion(self) -> None:
        heading = self.document.add_heading('Выводы по результатам испытаний', level=1)
        self._format_paragraph(heading)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Получаем данные о визуальном осмотре
        damage = self.data.get('damage_found', False)
        mount = self.data.get('mount_violation_found', False)
        weld = self.data.get('weld_violation_found', False)
        paint = self.data.get('paint_compliant', True)

        # Проверяем, есть ли критические проблемы (кроме покрытия)
        has_critical_problems = damage or mount or weld

        if has_critical_problems:
            # Если есть критические проблемы - не пригодны к эксплуатации
            text = "Конструкции маршевых лестниц не пригодны к эксплуатации."
            
            # Если защитное покрытие не соответствует - добавляем требование
            if not paint:
                text += " Требуется восстановить защитное покрытие."
        elif not paint:
            # Если только защитное покрытие не соответствует, а остальные поля в порядке
            text = (
                "Конструкции маршевых лестниц в прочностные испытания выдержали, "
                "нагрузка выдерживалась в течение 2 минут. После испытания не имеет трещин, "
                "прогибов, изломов. Требуется восстановить защитное покрытие. Соответствует ГОСТ Р 54253-2009 "
                "«Техника пожарная. Лестницы пожарные наружные стационарные. Ограждения кровли. "
                "Общие технические требования. Методы испытаний»."
            )
        else:
            # Если всё в порядке
            text = (
                "Конструкции маршевых лестниц в прочностные испытания выдержали, "
                "нагрузка выдерживалась в течение 2 минут. После испытания не имеет трещин, "
                "прогибов, изломов. Пригодны к эксплуатации. Соответствует ГОСТ Р 54253-2009 "
                "«Техника пожарная. Лестницы пожарные наружные стационарные. Ограждения кровли. "
                "Общие технические требования. Методы испытаний»."
            )

        paragraph = self.document.add_paragraph(text)
        self._format_paragraph(paragraph)
        self.document.add_paragraph()

    # --- Вспомогательные методы ------------------------------------------------

    @staticmethod
    def _to_float(value, default=0.0) -> float:
        try:
            return float(str(value).replace(',', '.'))
        except (TypeError, ValueError, AttributeError):
            return default


