"""
Базовые инструменты для генераторов протоколов
"""
from __future__ import annotations

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

import config
from logger import app_logger


class BaseProtocolGenerator(ABC):
    """Абстракция для всех генераторов протоколов"""

    def __init__(self, data: dict | None):
        self.data = data or {}
        self.document: Document | None = None
        config.ensure_directories()

    @abstractmethod
    def validate(self) -> None:
        """Должна выбрасывать ValueError при некорректных данных"""

    @abstractmethod
    def generate_doc(self, output_path: str | Path | None = None) -> str:
        """Создает протокол и возвращает путь к файлу"""

    # --- Общие служебные методы -------------------------------------------------

    def _require_fields(self, field_names: Iterable[str]) -> None:
        missing = [
            name for name in field_names
            if not str(self.data.get(name, "")).strip()
        ]
        if missing:
            raise ValueError(f"Не заполнены обязательные поля: {', '.join(missing)}")

    def _set_document(self) -> None:
        self.document = Document()
        self._setup_styles()

    def _setup_styles(self) -> None:
        if not self.document:
            return

        styles = self.document.styles
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = RGBColor(0, 0, 0)
        para = normal_style.paragraph_format
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.left_indent = Pt(0)
        para.right_indent = Pt(0)
        para.first_line_indent = None
        para.space_before = Pt(0)
        para.space_after = Pt(0)
        para.line_spacing = 1.0

        if 'CustomHeading' not in styles:
            style = styles.add_style('CustomHeading', 1)
            style.font.name = 'Times New Roman'
            style.font.size = Pt(11)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)

    def _format_paragraph(self, paragraph) -> None:  # noqa: ANN001
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

    def _add_company_header(self) -> None:
        if not self.document:
            return

        header_table = self.document.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.allow_autofit = False

        for row in header_table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(OxmlElement('w:tcBorders'))  # noqa: SLF001

        logo_cell = header_table.rows[0].cells[0]
        logo_cell.width = Inches(1.5)
        logo_cell.vertical_alignment = 1

        logo_file = config.get_logo_file()
        if logo_file and logo_file.exists():
            try:
                logo_paragraph = logo_cell.paragraphs[0]
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(str(logo_file.absolute()), height=Inches(0.8))
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                logo_paragraph.paragraph_format.space_before = Pt(0)
                logo_paragraph.paragraph_format.space_after = Pt(0)
            except Exception as exc:  # noqa: BLE001
                app_logger.warning(f"Не удалось добавить логотип: {exc}")
                logo_cell.text = config.COMPANY_NAME
        else:
            logo_cell.text = config.COMPANY_NAME

        details_cell = header_table.rows[0].cells[1]
        details_cell.width = Inches(4.0)
        details_cell.vertical_alignment = 1
        details_cell.text = ''

        lines = [
            (config.COMPANY_NAME, True),
            (config.COMPANY_ADDRESS_LINE1, False),
            (config.COMPANY_ADDRESS_LINE2, False),
            (config.COMPANY_PHONE, False),
            (config.COMPANY_EMAIL, False),
            (config.COMPANY_WEBSITE, False),
        ]

        for text, bold in lines:
            p = details_cell.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.bold = bold
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._format_paragraph(p)

        self.document.add_paragraph()  # пустая строка

    def _add_title(self, title: str, date_text: str | None = None) -> None:
        if not self.document:
            return
        heading = self.document.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = 1.0
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
        if date_text:
            date_para = self.document.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.paragraph_format.space_before = Pt(0)
            date_para.paragraph_format.space_after = Pt(0)
            date_run = date_para.add_run(date_text)
            date_run.font.bold = True
            date_run.font.size = Pt(11)
            self._format_paragraph(date_para)
        self.document.add_paragraph()

    def _add_key_value(self, title: str, value: str) -> None:
        if not self.document:
            return
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f"{title}: ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run = paragraph.add_run(value)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        self._format_paragraph(paragraph)

    def _add_table(self, headers: Sequence[str], rows: Sequence[Sequence[str]]):
        if not self.document:
            return None
        table = self.document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        for idx, header in enumerate(headers):
            cell = table.cell(0, idx)
            cell.text = header
            for paragraph in cell.paragraphs:
                self._format_paragraph(paragraph)
                paragraph.runs[0].font.bold = True if paragraph.runs else False
        for row_idx, row in enumerate(rows, start=1):
            for col_idx, value in enumerate(row):
                cell = table.cell(row_idx, col_idx)
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    self._format_paragraph(paragraph)
        self.document.add_paragraph()
        return table

    def _add_signatures(self) -> None:
        if not self.document:
            return
        self.document.add_paragraph()
        self.document.add_paragraph()

        paragraph = self.document.add_paragraph('Ответственный за проведение испытаний:')
        self._format_paragraph(paragraph)

        paragraph = self.document.add_paragraph('___________________ / _________________ /')
        self._format_paragraph(paragraph)

        paragraph = self.document.add_paragraph('        (подпись)                 (Ф.И.О.)')
        self._format_paragraph(paragraph)

    def _generate_filename(self, prefix: str) -> str:
        from datetime import datetime

        date_str = self.data.get('date', datetime.now().strftime('%Y-%m-%d'))
        date_str = date_str.replace('.', '-').replace('/', '-')
        time_str = datetime.now().strftime('%H-%M-%S')

        base_name = self.data.get('object_full_address') or prefix
        base_name = "".join(c for c in base_name if c.isalnum() or c in (' ', '-', '_')).strip()
        base_name = base_name[:50] if base_name else prefix

        return f"{prefix}_{date_str}_{time_str}_{base_name}.docx"

    def _resolve_output_path(self, output_path: str | Path | None, filename: str) -> Path:
        if output_path:
            path = Path(output_path)
            if path.is_dir():
                path = path / filename
        else:
            path = config.REPORTS_DIR / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        return path


