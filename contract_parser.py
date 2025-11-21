"""
Модуль для парсинга договоров и извлечения данных
"""
from docx import Document
from pathlib import Path
import re
from logger import app_logger


class ContractParser:
    """Класс для парсинга Word договоров"""
    
    def __init__(self, contracts_directory):
        """
        Args:
            contracts_directory (str or Path): Путь к папке с договорами
        """
        self.contracts_dir = Path(contracts_directory)
    
    def parse_contract(self, file_path):
        """
        Парсит один договор и извлекает данные
        
        Args:
            file_path (Path): Путь к файлу договора
        
        Returns:
            dict: Данные договора или None (если не удалось извлечь заказчика)
        """
        try:
            doc = Document(str(file_path))
            
            # Собираем весь текст документа
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Также проверяем таблицы
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            full_text.append(cell.text.strip())
            
            if not full_text:
                app_logger.warning(f"Файл {file_path.name} пустой")
                return None
            
            text = "\n".join(full_text)
            
            # Извлекаем данные
            customer = self._extract_customer(text)
            
            # Если не нашли заказчика - пропускаем договор
            if not customer:
                app_logger.warning(f"В файле {file_path.name} не найден заказчик")
                return None
            
            # Пытаемся извлечь данные из пункта 1.2
            section_12_data = self._extract_from_section_12(text)
            
            data = {
                'file_name': file_path.name,
                'file_path': str(file_path),
                'customer': customer,
                'object_full_address': section_12_data.get('combined_address'),  # Объединённое поле
            }
            
            # Логируем что нашли
            found_fields = [f"заказчик: {customer}"]
            if data['object_full_address']:
                preview = data['object_full_address'][:60] + '...' if len(data['object_full_address']) > 60 else data['object_full_address']
                found_fields.append(f"адрес/наименование: {preview}")
            
            app_logger.info(f"Распарсен {file_path.name}: {', '.join(found_fields)}")
            return data
            
        except Exception as e:
            app_logger.error(f"Ошибка при парсинге {file_path.name}: {e}")
            return None
    
    def _extract_from_section_12(self, text):
        """
        Специальная обработка пункта 1.2 договора.
        Извлекает весь текст после "на объекте заказчика" до конца пункта.
        
        Пример:
        1.2. ...на объекте заказчика: «Многоэтажный дом...» по адресу: г. Екатеринбург...
        
        Извлекает всё от "на объекте заказчика" до конца пункта (включая наименование и адрес).
        
        Returns:
            dict: {'combined_address': str or None}
        """
        result = {'combined_address': None}
        
        try:
            # Ищем пункт 1.2 целиком (захватываем до следующего пункта)
            section_pattern = r'1\.2\s*\.?\s*(.*?)(?=\n\s*1\.[3-9]|\n\s*2\.|\Z)'
            section_match = re.search(section_pattern, text, re.IGNORECASE | re.DOTALL)
            
            if not section_match:
                # Запасной вариант
                section_pattern = r'1\.2\s*\.?\s*([^\n]*(?:\n(?!\s*1\.[3-9]|\s*2\.)[^\n]*){0,20})'
                section_match = re.search(section_pattern, text, re.IGNORECASE | re.DOTALL)
            
            if not section_match:
                return result
            
            section_text = section_match.group(1)
            
            # Извлекаем весь текст после "на объекте заказчика"
            object_patterns = [
                # После "на объекте заказчика:" - берём всё до конца пункта
                r'на\s+объекте\s+заказчика[:\s]*[-–—]?\s*(.+)',
                # Альтернатива
                r'объект[е]?\s+заказчика[:\s]*[-–—]?\s*(.+)',
            ]
            
            for pattern in object_patterns:
                match = re.search(pattern, section_text, re.IGNORECASE | re.DOTALL)
                if match:
                    combined = match.group(1).strip()
                    
                    # Убираем переносы строк, заменяем на пробелы
                    combined = ' '.join(combined.split())
                    
                    # Убираем возможный мусор в конце (следующие предложения)
                    combined = re.sub(r'\s+(?:Исполнитель|Срок\s+выполнения|В\s+течени[ие]).*$', '', combined, flags=re.IGNORECASE)
                    
                    # Очистка от лишних пробелов и символов
                    combined = combined.strip()
                    
                    if len(combined) > 5 and len(combined) < 1000:
                        result['combined_address'] = combined
                        break
            
        except Exception as e:
            app_logger.debug(f"Ошибка при извлечении из пункта 1.2: {e}")
        
        return result
    
    def _extract_customer(self, text):
        """
        Извлекает название заказчика из преамбулы договора.
        Формат: "Общество с ограниченной ответственностью «НАЗВАНИЕ», именуемое в дальнейшем «Заказчик»"
        """
        # Основной паттерн - полное название организации до "именуемое"
        patterns = [
            # ООО с кавычками
            r'(Общество\s+с\s+ограниченной\s+ответственностью\s*["\s\'«]+[^"\'»]+["\s\'»]+)\s*,?\s*именуем[а-я]{0,3}\s+в\s+дальнейшем\s*["\s\'«]*\s*[Зз]аказчик',
            # Краткая форма ООО
            r'(ООО\s*["\s\'«]+[^"\'»]+["\s\'»]+)\s*,?\s*именуем[а-я]{0,3}\s+в\s+дальнейшем\s*["\s\'«]*\s*[Зз]аказчик',
            # АО, ЗАО, ПАО
            r'((?:Закрытое|Открытое|Публичное)?\s*[Аа]кционерное\s+общество\s*["\s\'«]+[^"\'»]+["\s\'»]+)\s*,?\s*именуем[а-я]{0,3}\s+в\s+дальнейшем\s*["\s\'«]*\s*[Зз]аказчик',
            r'([ЗОП]АО\s*["\s\'«]+[^"\'»]+["\s\'»]+)\s*,?\s*именуем[а-я]{0,3}\s+в\s+дальнейшем\s*["\s\'«]*\s*[Зз]аказчик',
            # ИП
            r'(Индивидуальный\s+предприниматель\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)\s*,?\s*именуем[а-я]{0,3}\s+в\s+дальнейшем\s*["\s\'«]*\s*[Зз]аказчик',
            r'(ИП\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)\s*,?\s*именуем[а-я]{0,3}\s+в\s+дальнейшем\s*["\s\'«]*\s*[Зз]аказчик',
            # Общий паттерн - любой текст перед "именуемое"
            r',\s*([^,]{10,200}?)\s*,?\s*именуем[а-я]{0,3}\s+в\s+дальнейшем\s*["\s\'«]*\s*[Зз]аказчик',
        ]
        
        for pattern in patterns:
            try:
                match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                if match:
                    customer = match.group(1).strip()
                    # Очистка от лишних символов
                    customer = customer.strip('",\'«»:;.\n\r\t ')
                    # Убираем служебные слова если попали
                    customer = re.sub(r'^(в\s+лице|далее)\s*[-–—]?\s*', '', customer, flags=re.IGNORECASE)
                    
                    if len(customer) > 5 and len(customer) < 250:
                        return customer
            except Exception:
                continue
        
        return None
    
    def _extract_object_name(self, text):
        """
        Извлекает наименование объекта из пункта 1.2 договора.
        Ищет после слова "наименование".
        """
        patterns = [
            # Ищем в пункте 1.2 после слова "наименование"
            r'1\.2[.\s]*[^\n]*?[Нн]аименование[:\s]+([^\n;\.]+)',
            r'1\.2[.\s]*[^\n]*?[Нн]аименование[:\s]*[-–—]?\s*([^\n;\.]+)',
            # Запасные варианты
            r'[Нн]аименование\s+объекта[:\s]+([^\n;\.]+)',
            r'[Нн]аименование[:\s]+([^\n;\.]{10,})',
            r'[Оо]бъект[:\s]+([^\n;\.]+)',
            r'[Мм]есто проведения работ[:\s]+([^\n;\.]+)',
            r'[Оо]бъект испытания[:\s]+([^\n;\.]+)',
        ]
        
        for pattern in patterns:
            try:
                match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                if match:
                    obj_name = match.group(1).strip()
                    # Очистка
                    obj_name = obj_name.strip('",\'«»:;.\n\r\t ')
                    # Удаляем возможные слова "адрес" в конце
                    obj_name = re.sub(r'\s*[Аа]дрес.*$', '', obj_name)
                    
                    if len(obj_name) > 3 and len(obj_name) < 300:
                        return obj_name
            except Exception:
                continue
        
        return None
    
    def _extract_address(self, text):
        """
        Извлекает адрес объекта из пункта 1.2 договора.
        Ищет после слова "адрес".
        """
        patterns = [
            # Ищем в пункте 1.2 после слова "адрес"
            r'1\.2[.\s]*.*?[Аа]дрес[:\s]+([^\n;\.]+(?:г\.|город|обл\.|область|ул\.|улица|д\.|дом|пр\.|проспект)[^\n;\.]+)',
            r'1\.2[.\s]*.*?[Аа]дрес[:\s]*[-–—]?\s*([^\n;\.]+)',
            # Общий поиск адреса объекта
            r'[Аа]дрес\s+объекта[:\s]+([^\n;\.]+)',
            r'[Аа]дрес[:\s]+([^\n;]+(?:г\.|город|обл\.|область|ул\.|улица|д\.|дом)[^\n;]+)',
            # Поиск по началу с "г." или "город"
            r'(?:г\.|город)\s*([А-ЯЁ][а-яё]+[,\s]+[^\n]{10,150})',
        ]
        
        for pattern in patterns:
            try:
                match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                if match:
                    address = match.group(1).strip()
                    # Очистка
                    address = address.strip('",\'«»:;.\n\r\t ')
                    # Убираем возможный мусор в конце
                    address = re.sub(r'\s*(далее|именуемый|именуемое).*$', '', address, flags=re.IGNORECASE)
                    
                    if len(address) > 5 and len(address) < 300:
                        return address
            except Exception:
                continue
        
        return None
    
    def scan_contracts_directory(self):
        """
        Сканирует папку с договорами и извлекает данные из всех файлов
        
        Returns:
            list: Список словарей с данными договоров
        """
        if not self.contracts_dir.exists():
            app_logger.warning(f"Папка с договорами не найдена: {self.contracts_dir}")
            return []
        
        contracts_data = []
        total_files = 0
        
        # Ищем все .docx файлы
        for file_path in self.contracts_dir.glob("*.docx"):
            # Пропускаем временные файлы Word
            if file_path.name.startswith('~$'):
                continue
            
            total_files += 1
            
            try:
                data = self.parse_contract(file_path)
                if data and data.get('customer'):  # Добавляем только если нашли заказчика
                    contracts_data.append(data)
            except Exception as e:
                app_logger.error(f"Ошибка обработки {file_path.name}: {e}")
                continue
        
        app_logger.info(f"Обработано файлов: {total_files}, успешно извлечено: {len(contracts_data)}")
        return contracts_data

